from typing import Dict, Any, List
from docx import Document
from .base_processor import BaseProcessor
from .registry import ProcessorRegistry
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os


@ProcessorRegistry.register
class DocxProcessor(BaseProcessor):
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        return ['.docx']
    
    def process(self, file_path: str) -> List[Dict[str, Any]]:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join([cell.text for cell in row.cells])
        
        chunks = self.text_splitter.split_text(text)
        return [{'content': chunk, 'metadata': self.get_metadata(file_path)} for chunk in chunks]
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)
        return {
            'file_type': 'docx',
            'file_name': os.path.basename(file_path),
            'file_size': os.path.getsize(file_path),
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'core_properties': {
                'author': doc.core_properties.author,
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
                'title': doc.core_properties.title,
                'subject': doc.core_properties.subject
            }
        } 